"""
Modul pentru extragerea datelor din fișa disciplinei în format DOCX.
Folosește indexare directă pentru acces rapid la celule.
"""
import re
from docx import Document
from typing import Dict, Optional


def clean_text(text: str) -> str:
    """Curăță textul de spații multiple și caractere speciale."""
    text = text.strip()
    text = re.sub(r'\s+', ' ', text)
    return text


def extract_fisa_disciplina(file_path: str) -> Dict[str, any]:
    """
    Extrage datele din fișa disciplinei folosind indexare directă.
    
    Structura fișei (indexare 0-based):
    - Tabel 1 (index 1): Date despre disciplină
      - Rând 0, Celula 4: Denumiri (română / engleză)
      - Rând 1, Celula 4: Cod disciplină
      - Rând 1, Celula 8: Categoria (DA/DOP/DOB/DFA)
      - Rând 4, Celula 6: Evaluare (E/V/C)
    
    - Tabel 2 (index 2): Ore pe săptămână și distribuție timp
      - Rând 0, Celula 4: Ore curs
      - Rând 0, Celula 8: Ore laborator
      - Rând 3, Celula 12: Ore studiu manual
      - Rând 4, Celula 12: Ore documentare
      - Rând 5, Celula 12: Ore pregătire seminarii
      - Rând 6, Celula 12: Ore examinări
      - Rând 10, Celula 1: Număr credite
    
    Args:
        file_path: Calea către fișierul DOCX
        
    Returns:
        Dicționar cu datele extrase
    """
    doc = Document(file_path)
    
    result = {
        'cod': None,
        'denumire_ro': None,
        'denumire_en': None,
        'categoria': None,
        'credite': None,
        'nr_ore_saptamana_total': None,
        'nr_ore_saptamana': {
            'curs': 0,
            'seminar': 0,
            'proiect': 0,
            'lucrari': 0
        },
        'total_ore_plan': None,
        'distributie_fond_timp': {
            'studiu_manual': 0,
            'documentare': 0,
            'pregatire_seminarii': 0,
            'examinari': 0
        },
        'total_ore_studiu_individual': None,
        'total_ore_semestru': None,
        'evaluare': None
    }
    
    try:
        # Tabelul 2 (index 1) - Date despre disciplină
        table_disciplina = doc.tables[1]
        
        # Denumiri - Rând 0, Celula 4
        denumiri_text = clean_text(table_disciplina.rows[0].cells[4].text)
        if '/' in denumiri_text:
            parts = denumiri_text.split('/')
            result['denumire_ro'] = clean_text(parts[0])
            result['denumire_en'] = clean_text(parts[1]) if len(parts) > 1 else None
        
        # Cod - Rând 1, Celula 4
        cod_text = clean_text(table_disciplina.rows[1].cells[4].text)
        match = re.search(r'([A-Z]{2}\.[A-Z]{2}\.\d{3})', cod_text)
        if match:
            result['cod'] = match.group(1)
        
        # Categoria - Rând 1, Celula 8
        categoria_text = clean_text(table_disciplina.rows[1].cells[8].text)
        match = re.search(r'\b(DA|DOP|DOB|DFA)\b', categoria_text)
        if match:
            result['categoria'] = match.group(1)
        
        # Evaluare - Rând 4, Celula 6
        evaluare_text = clean_text(table_disciplina.rows[4].cells[6].text)
        match = re.search(r'\b([EVC])\b', evaluare_text)
        if match:
            result['evaluare'] = match.group(1)
        
        # Tabelul 3 (index 2) - Ore pe săptămână și totalizări
        table_ore = doc.tables[2]
        
        # Total ore pe săptămână - Rând 0, Celula 1
        total_ore_saptamana_text = clean_text(table_ore.rows[0].cells[1].text)
        match = re.search(r'\b(\d+)\b', total_ore_saptamana_text)
        if match:
            result['nr_ore_saptamana_total'] = int(match.group(1))
        
        # Ore curs - Rând 0, Celula 4
        ore_curs_text = clean_text(table_ore.rows[0].cells[4].text)
        match = re.search(r'\b(\d+)\b', ore_curs_text)
        if match:
            result['nr_ore_saptamana']['curs'] = int(match.group(1))
        
        # Ore laborator - Rând 0, Celula 8
        ore_lab_text = clean_text(table_ore.rows[0].cells[8].text)
        match = re.search(r'\b(\d+)\b', ore_lab_text)
        if match:
            result['nr_ore_saptamana']['lucrari'] = int(match.group(1))
        
        # Seminar și proiect rămân 0 (nu sunt în structura actuală)
        
        # Total ore din planul de învățământ - Rând 1, Celula 1
        total_ore_plan_text = clean_text(table_ore.rows[1].cells[1].text)
        match = re.search(r'\b(\d+)\b', total_ore_plan_text)
        if match:
            result['total_ore_plan'] = int(match.group(1))
        
        # Credite - Rând 10, Celula 1
        credite_text = clean_text(table_ore.rows[10].cells[1].text)
        match = re.search(r'\b(\d+)\b', credite_text)
        if match:
            result['credite'] = int(match.group(1))
        
        # Distribuția fondului de timp - Rândurile 3-6, Celula 12
        # Studiul după manual - Rând 3, Celula 12
        studiu_text = clean_text(table_ore.rows[3].cells[12].text)
        match = re.search(r'\b(\d+)\b', studiu_text)
        if match:
            result['distributie_fond_timp']['studiu_manual'] = int(match.group(1))
        
        # Documentare - Rând 4, Celula 12
        doc_text = clean_text(table_ore.rows[4].cells[12].text)
        match = re.search(r'\b(\d+)\b', doc_text)
        if match:
            result['distributie_fond_timp']['documentare'] = int(match.group(1))
        
        # Pregătire seminarii - Rând 5, Celula 12
        prep_text = clean_text(table_ore.rows[5].cells[12].text)
        match = re.search(r'\b(\d+)\b', prep_text)
        if match:
            result['distributie_fond_timp']['pregatire_seminarii'] = int(match.group(1))
        
        # Examinări - Rând 6, Celula 12
        exam_text = clean_text(table_ore.rows[6].cells[12].text)
        match = re.search(r'\b(\d+)\b', exam_text)
        if match:
            result['distributie_fond_timp']['examinari'] = int(match.group(1))
        
        # Total ore studiu individual - Rând 8, Celula 1
        total_studiu_text = clean_text(table_ore.rows[8].cells[1].text)
        match = re.search(r'\b(\d+)\b', total_studiu_text)
        if match:
            result['total_ore_studiu_individual'] = int(match.group(1))
        
        # Total ore pe semestru - Rând 9, Celula 1
        total_semestru_text = clean_text(table_ore.rows[9].cells[1].text)
        match = re.search(r'\b(\d+)\b', total_semestru_text)
        if match:
            result['total_ore_semestru'] = int(match.group(1))
        
    except IndexError as e:
        print(f"Eroare la accesarea structurii: {e}")
    
    return result

if __name__ == '__main__':
    # Test pe fișa încărcată
    import json
    
    fisa_path = '/mnt/user-data/uploads/IGIA202_Modelare_in__ingineria_geotehnica_Teodoru_2.docx'
    
    print("Extrag datele din fișa disciplinei...")
    date = extract_fisa_disciplina(fisa_path)
    
    print("\n=== DATE EXTRASE ===")
    print(json.dumps(date, indent=2, ensure_ascii=False))